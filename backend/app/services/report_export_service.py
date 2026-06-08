"""Export scientific reports to Markdown, Word, and PDF with GxP traceability."""

from __future__ import annotations

import io
import re
from datetime import datetime

from app.models.db_models import ScientificReport
from app.services.report_content_builder import enrich_report_content

REPORT_TYPE_LABELS = {
    "hypothesis_report": "Hypothesis Report",
    "experiment_plan": "Experiment Plan",
    "study_report": "Study Report",
    "target_discovery": "Target Discovery Report",
    "cmc_readiness": "CMC Readiness Assessment",
    "meeting_brief": "Meeting Brief",
}


def _safe_filename(title: str, ext: str) -> str:
    safe = re.sub(r"[^\w\s-]", "", title or "report").strip().replace(" ", "_")[:60]
    return f"{safe or 'report'}.{ext}"


def _escape_pdf(text: str) -> str:
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _report_payload(report: ScientificReport) -> dict:
    c = dict(report.content_json or {})
    sections = c.get("section_content") or []
    if not sections and c.get("sections"):
        raw = c["sections"]
        sections = [
            {"name": s, "content": ""} if isinstance(s, str) else s
            for s in raw
        ]

    body = c.get("body") or c.get("answer") or ""
    sparse = (
        not any(isinstance(s, dict) and (s.get("content") or "").strip() for s in sections)
        and len(body) < 400
    )
    if sparse:
        source = c.get("generated_from") or {}
        enriched = enrich_report_content(
            report.report_type,
            c,
            query=str(source.get("query") or source.get("project") or report.title),
            citations=c.get("citations_list") or [],
        )
        c = {**c, **enriched}
        sections = c.get("section_content") or sections
        body = c.get("body") or body

    if not sections and body:
        sections = [{"name": "Full Report", "content": body}]

    return {
        "title": report.title,
        "report_type": report.report_type,
        "report_type_label": REPORT_TYPE_LABELS.get(
            report.report_type, report.report_type.replace("_", " ").title(),
        ),
        "status": report.status,
        "summary": c.get("summary") or "",
        "body": body,
        "sections": sections,
        "citations": c.get("citations_list") or [],
        "limitations": c.get("limitations") or [],
        "recommendations": c.get("recommendations") or [],
        "figures": c.get("figures_list") or c.get("figures") or [],
        "hypotheses": c.get("hypotheses") or [],
        "experiment_plan": c.get("experiment_plan"),
        "doe_design": c.get("doe_design"),
        "gaps": c.get("gaps") or [],
        "verdict": c.get("verdict"),
        "evidence_for": c.get("evidence_for") or [],
        "evidence_against": c.get("evidence_against") or [],
        "findings": c.get("findings") or [],
        "confidence": c.get("confidence"),
        "model_used": c.get("model_used"),
        "mode": c.get("mode"),
        "word_count": c.get("word_count", 0),
        "generated_at": report.created_at.strftime("%Y-%m-%d %H:%M UTC") if report.created_at else "",
        "workflow_run_id": report.workflow_run_id,
    }


def _gxp_header_lines(p: dict, report_id: str) -> list[str]:
    return [
        f"# {p['title']}",
        "",
        "| Field | Value |",
        "| --- | --- |",
        f"| Report ID | `{report_id}` |",
        f"| Type | {p['report_type_label']} |",
        f"| Status | {p['status']} |",
        f"| Generated | {p['generated_at']} |",
        f"| Word count | {p.get('word_count') or 'n/a'} |",
        f"| Mode | {p.get('mode') or 'n/a'} |",
        f"| Model | {p.get('model_used') or 'n/a'} |",
        f"| Confidence | {p.get('confidence') if p.get('confidence') is not None else 'n/a'} |",
        "",
    ]


def _hypothesis_md(h: dict, index: int) -> list[str]:
    lines = [f"### {index}. {h.get('title', h.get('hypothesis', 'Hypothesis'))}"]
    if h.get("statement"):
        lines.append(f"**Statement:** {h['statement']}")
    for key, label in [
        ("target_or_mechanism", "Target / mechanism"),
        ("disease_context", "Disease context"),
        ("rationale", "Rationale"),
        ("evidence_strength", "Evidence strength"),
    ]:
        if h.get(key):
            lines.append(f"**{label}:** {h[key]}")
    if h.get("confidence") is not None:
        lines.append(f"**Confidence:** {h['confidence']}")
    if h.get("suggested_experiments"):
        lines.append("**Suggested experiments:**")
        lines.extend(f"- {x}" for x in h["suggested_experiments"])
    lines.append("")
    return lines


def _export_sections_markdown(p: dict) -> list[str]:
    lines: list[str] = []
    for section in p["sections"]:
        if isinstance(section, dict):
            name = section.get("name") or "Section"
            content = (section.get("content") or "").strip()
            if content:
                lines += [f"## {name}", content, ""]
        elif isinstance(section, str) and section.strip():
            lines.append(f"## {section}")
            lines.append("")
    return lines


def export_report_markdown(report: ScientificReport) -> bytes:
    p = _report_payload(report)
    lines = _gxp_header_lines(p, report.id)

    if p["summary"] and not any(
        s.get("name") == "Executive Summary" for s in p["sections"] if isinstance(s, dict)
    ):
        lines += ["## Executive Summary", p["summary"], ""]

    lines += _export_sections_markdown(p)

    if p["hypotheses"] and not any(
        "hypoth" in (s.get("name") or "").lower() for s in p["sections"] if isinstance(s, dict)
    ):
        lines += ["## Hypotheses"]
        for i, h in enumerate(p["hypotheses"], 1):
            if isinstance(h, dict):
                lines.extend(_hypothesis_md(h, i))
            else:
                lines.append(f"- {h}")
        lines.append("")

    if p["verdict"]:
        lines += ["## Validation Verdict", str(p["verdict"]).replace("_", " ").title(), ""]

    if p["limitations"]:
        lines += ["## Limitations", *[f"- {x}" for x in p["limitations"]], ""]
    if p["recommendations"]:
        lines += ["## Recommendations", *[f"- {x}" for x in p["recommendations"]], ""]

    if p["citations"] and not any(
        "reference" in (s.get("name") or "").lower() for s in p["sections"] if isinstance(s, dict)
    ):
        lines += ["## References & Evidence"]
        for cite in p["citations"]:
            idx = cite.get("index", "?")
            title = cite.get("title", "Source")
            source = cite.get("source", "")
            excerpt = (cite.get("excerpt") or "")[:400]
            url = cite.get("url") or ""
            line = f"[{idx}] **{title}** ({source})"
            if excerpt:
                line += f"\n> {excerpt}"
            if url:
                line += f"\n> {url}"
            lines.append(line)
        lines.append("")

    lines += [
        "---",
        f"*SciNova OS Scientific Report · {p['report_type_label']} · {datetime.utcnow().strftime('%Y-%m-%d')}*",
    ]
    return "\n".join(lines).encode("utf-8")


def export_report_docx(report: ScientificReport) -> bytes:
    from docx import Document as DocxDocument

    p = _report_payload(report)
    doc = DocxDocument()
    doc.add_heading(p["title"], 0)
    doc.add_paragraph(
        f"Type: {p['report_type_label']} · Status: {p['status']} · ID: {report.id}\n"
        f"Generated: {p['generated_at']} · Words: {p.get('word_count', 'n/a')} · "
        f"Confidence: {p.get('confidence', 'n/a')}"
    )

    for section in p["sections"]:
        if isinstance(section, dict):
            name = section.get("name") or "Section"
            content = (section.get("content") or "").strip()
            if content:
                doc.add_heading(name, level=1)
                for para in content.split("\n\n"):
                    text = para.strip()
                    if text:
                        doc.add_paragraph(text)

    if p["limitations"]:
        doc.add_heading("Limitations", level=1)
        for item in p["limitations"]:
            doc.add_paragraph(str(item), style="List Bullet")

    if p["recommendations"]:
        doc.add_heading("Recommendations", level=1)
        for item in p["recommendations"]:
            doc.add_paragraph(str(item), style="List Bullet")

    if p["citations"]:
        doc.add_heading("References & Evidence", level=1)
        for cite in p["citations"]:
            idx = cite.get("index", "?")
            excerpt = (cite.get("excerpt") or "")[:300]
            doc.add_paragraph(
                f"[{idx}] {cite.get('title', 'Source')} ({cite.get('source', '')})",
                style="List Bullet",
            )
            if excerpt:
                doc.add_paragraph(excerpt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_report_pdf(report: ScientificReport) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    p = _report_payload(report)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "ReportBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "ReportH2",
        parent=styles["Heading2"],
        fontSize=13,
        leading=16,
        spaceBefore=12,
        spaceAfter=6,
        textColor="#0e7490",
    )

    story = [
        Paragraph(_escape_pdf(p["title"]), styles["Title"]),
        Spacer(1, 8),
        Paragraph(
            _escape_pdf(
                f"Type: {p['report_type_label']} · Status: {p['status']} · "
                f"Generated: {p['generated_at']} · Words: {p.get('word_count', 'n/a')} · "
                f"Confidence: {p.get('confidence', 'n/a')}"
            ),
            body_style,
        ),
        Spacer(1, 14),
    ]

    section_count = 0
    for section in p["sections"]:
        if not isinstance(section, dict):
            continue
        content = (section.get("content") or "").strip()
        if not content:
            continue
        name = section.get("name") or "Section"
        story.append(Paragraph(_escape_pdf(name), h2_style))
        for para in content.split("\n\n"):
            text = para.strip()
            if text:
                story.append(Paragraph(_escape_pdf(text), body_style))
        story.append(Spacer(1, 8))
        section_count += 1
        if section_count > 0 and section_count % 6 == 0:
            story.append(PageBreak())

    if not section_count and p["body"]:
        story.append(Paragraph("Full Report", h2_style))
        for para in p["body"].split("\n\n"):
            text = para.strip()
            if text:
                story.append(Paragraph(_escape_pdf(text), body_style))

    if p["citations"] and not any(
        "reference" in (s.get("name") or "").lower() for s in p["sections"] if isinstance(s, dict)
    ):
        story.append(Spacer(1, 12))
        story.append(Paragraph("References &amp; Evidence", h2_style))
        for cite in p["citations"]:
            idx = cite.get("index", "?")
            title = cite.get("title", "Source")
            source = cite.get("source", "")
            excerpt = (cite.get("excerpt") or "")[:350]
            line = f"[{idx}] {title} ({source})"
            if excerpt:
                line += f"<br/>{excerpt}"
            story.append(Paragraph(_escape_pdf(line), body_style))

    if len(story) <= 3:
        story.append(Paragraph(
            "This report has limited stored content. Regenerate the report to produce a full export.",
            body_style,
        ))

    doc.build(story)
    return buf.getvalue()


def export_report(report: ScientificReport, fmt: str) -> tuple[bytes, str, str]:
    fmt = fmt.lower()
    if fmt in ("md", "markdown"):
        return (
            export_report_markdown(report),
            _safe_filename(report.title, "md"),
            "text/markdown; charset=utf-8",
        )
    if fmt == "docx":
        return (
            export_report_docx(report),
            _safe_filename(report.title, "docx"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if fmt == "pdf":
        return (
            export_report_pdf(report),
            _safe_filename(report.title, "pdf"),
            "application/pdf",
        )
    raise ValueError(f"Unsupported format: {fmt}")
