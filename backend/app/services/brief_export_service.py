"""Export meeting briefs to Markdown, Word, and PDF."""

from __future__ import annotations

import io
from datetime import datetime

from app.models.db_models import ScientificReport


def _brief_sections(report: ScientificReport) -> dict:
    c = report.content_json or {}
    return {
        "title": report.title,
        "summary": c.get("summary") or "",
        "body": c.get("body") or c.get("summary") or "",
        "agenda": c.get("agenda") or [],
        "key_findings": c.get("key_findings") or [],
        "decisions_needed": c.get("decisions_needed") or [],
        "action_items": c.get("action_items") or [],
        "risks": c.get("risks") or [],
        "topic": c.get("topic") or "",
        "audience": c.get("audience") or "",
    }


def export_brief_markdown(report: ScientificReport) -> bytes:
    s = _brief_sections(report)
    lines = [
        f"# {s['title']}",
        "",
        f"*Generated {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} · Audience: {s['audience']}*",
        "",
        "## Executive Summary",
        s["summary"],
        "",
    ]
    if s["agenda"]:
        lines += ["## Agenda", *[f"- {a}" for a in s["agenda"]], ""]
    if s["key_findings"]:
        lines += ["## Key Findings", *[f"- {f}" for f in s["key_findings"]], ""]
    if s["decisions_needed"]:
        lines += ["## Decisions Needed", *[f"- {d}" for d in s["decisions_needed"]], ""]
    if s["action_items"]:
        lines += ["## Action Items"]
        for item in s["action_items"]:
            if isinstance(item, dict):
                lines.append(f"- **{item.get('owner', 'TBD')}**: {item.get('task', '')} ({item.get('due_hint', '')})")
            else:
                lines.append(f"- {item}")
        lines.append("")
    if s["risks"]:
        lines += ["## Risks", *[f"- {r}" for r in s["risks"]], ""]
    if s["body"]:
        lines += ["## Full Brief", s["body"], ""]
    return "\n".join(lines).encode("utf-8")


def export_brief_docx(report: ScientificReport) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    s = _brief_sections(report)
    doc = DocxDocument()
    doc.add_heading(s["title"], 0)
    doc.add_paragraph(f"Audience: {s['audience']} · {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(s["summary"])

    if s["agenda"]:
        doc.add_heading("Agenda", level=1)
        for item in s["agenda"]:
            doc.add_paragraph(item, style="List Bullet")

    if s["key_findings"]:
        doc.add_heading("Key Findings", level=1)
        for item in s["key_findings"]:
            doc.add_paragraph(item, style="List Bullet")

    if s["decisions_needed"]:
        doc.add_heading("Decisions Needed", level=1)
        for item in s["decisions_needed"]:
            doc.add_paragraph(item, style="List Bullet")

    if s["action_items"]:
        doc.add_heading("Action Items", level=1)
        for item in s["action_items"]:
            if isinstance(item, dict):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{item.get('owner', 'TBD')}: ").bold = True
                p.add_run(f"{item.get('task', '')} ({item.get('due_hint', '')})")
            else:
                doc.add_paragraph(str(item), style="List Bullet")

    if s["body"]:
        doc.add_heading("Full Brief", level=1)
        for para in s["body"].split("\n\n"):
            doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_brief_pdf(report: ScientificReport) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    s = _brief_sections(report)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(s["title"], styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"<b>Audience:</b> {s['audience']}", styles["Normal"]),
        Spacer(1, 12),
        Paragraph("<b>Executive Summary</b>", styles["Heading2"]),
        Paragraph(s["summary"].replace("\n", "<br/>"), styles["Normal"]),
    ]

    def add_bullets(heading: str, items: list) -> None:
        if not items:
            return
        story.append(Spacer(1, 8))
        story.append(Paragraph(f"<b>{heading}</b>", styles["Heading2"]))
        for item in items:
            text = item if isinstance(item, str) else f"{item.get('owner')}: {item.get('task')}"
            story.append(Paragraph(f"• {text}", styles["Normal"]))

    add_bullets("Agenda", s["agenda"])
    add_bullets("Key Findings", s["key_findings"])
    add_bullets("Decisions Needed", s["decisions_needed"])
    add_bullets("Action Items", s["action_items"])
    add_bullets("Risks", s["risks"])

    if s["body"]:
        story.append(Spacer(1, 8))
        story.append(Paragraph("<b>Full Brief</b>", styles["Heading2"]))
        story.append(Paragraph(s["body"].replace("\n", "<br/>")[:8000], styles["Normal"]))

    doc.build(story)
    return buf.getvalue()


def export_brief(report: ScientificReport, fmt: str) -> tuple[bytes, str, str]:
    fmt = fmt.lower()
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in report.title[:60])
    if fmt == "docx":
        return export_brief_docx(report), f"{safe}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        return export_brief_pdf(report), f"{safe}.pdf", "application/pdf"
    return export_brief_markdown(report), f"{safe}.md", "text/markdown; charset=utf-8"
