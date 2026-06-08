"""Document parsing for Scientific Data Fabric."""

import json
import csv
import io
from pathlib import Path
from dataclasses import dataclass, field


@dataclass
class ParsedDocument:
    text: str
    metadata: dict = field(default_factory=dict)
    sections: list[dict] = field(default_factory=list)


def parse_document(file_path: str, file_format: str) -> ParsedDocument:
    fmt = file_format.upper().lstrip(".")
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    parsers = {
        "PDF": _parse_pdf,
        "DOCX": _parse_docx,
        "CSV": _parse_csv,
        "XLSX": _parse_xlsx,
        "JSON": _parse_json,
        "TXT": _parse_txt,
    }
    parser = parsers.get(fmt, _parse_txt)
    return parser(path)


def _parse_pdf(path: Path) -> ParsedDocument:
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append({"page": i + 1, "text": text.strip()})

    full_text = "\n\n".join(p["text"] for p in pages)
    meta = doc.metadata or {}
    return ParsedDocument(
        text=full_text,
        metadata={
            "pages": len(doc),
            "author": meta.get("author", ""),
            "title": meta.get("title", ""),
            "subject": meta.get("subject", ""),
            "file_size_bytes": path.stat().st_size,
        },
        sections=[{"type": "page", **p} for p in pages],
    )


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    return ParsedDocument(
        text=full_text,
        metadata={
            "paragraphs": len(paragraphs),
            "file_size_bytes": path.stat().st_size,
        },
        sections=[{"type": "paragraph", "index": i, "text": p} for i, p in enumerate(paragraphs)],
    )


def _parse_csv(path: Path) -> ParsedDocument:
    import pandas as pd

    df = pd.read_csv(str(path))
    rows = []
    for i, row in df.iterrows():
        row_text = " | ".join(f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col]))
        rows.append({"row": int(i), "text": row_text})

    summary = f"CSV dataset with {len(df)} rows and columns: {', '.join(df.columns.astype(str))}"
    full_text = summary + "\n\n" + "\n".join(r["text"] for r in rows[:500])
    return ParsedDocument(
        text=full_text,
        metadata={
            "rows": len(df),
            "columns": list(df.columns.astype(str)),
            "file_size_bytes": path.stat().st_size,
        },
        sections=[{"type": "row", **r} for r in rows],
    )


def _parse_xlsx(path: Path) -> ParsedDocument:
    import pandas as pd

    xls = pd.ExcelFile(str(path))
    sections = []
    all_text_parts = [f"Excel workbook with sheets: {', '.join(xls.sheet_names)}"]

    for sheet in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet)
        all_text_parts.append(f"\n[Sheet: {sheet}] — {len(df)} rows")
        for i, row in df.head(200).iterrows():
            row_text = " | ".join(f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col]))
            sections.append({"type": "row", "sheet": sheet, "row": int(i), "text": row_text})
            all_text_parts.append(row_text)

    return ParsedDocument(
        text="\n".join(all_text_parts),
        metadata={
            "sheets": xls.sheet_names,
            "file_size_bytes": path.stat().st_size,
        },
        sections=sections,
    )


def _parse_json(path: Path) -> ParsedDocument:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    def flatten(obj, prefix=""):
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                lines.extend(flatten(v, f"{prefix}{k}."))
        elif isinstance(obj, list):
            for i, item in enumerate(obj[:200]):
                lines.extend(flatten(item, f"{prefix}[{i}]."))
        else:
            lines.append(f"{prefix.rstrip('.')}: {obj}")
        return lines

    lines = flatten(data)
    full_text = "\n".join(lines)
    return ParsedDocument(
        text=full_text,
        metadata={
            "json_type": type(data).__name__,
            "file_size_bytes": path.stat().st_size,
        },
        sections=[{"type": "field", "text": line} for line in lines[:500]],
    )


def _parse_txt(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(
        text=text,
        metadata={"file_size_bytes": path.stat().st_size, "char_count": len(text)},
        sections=[{"type": "full", "text": text}],
    )


def infer_source_type(filename: str, declared: str | None = None) -> str:
    if declared and declared != "file_upload":
        return declared
    ext = Path(filename).suffix.lower()
    mapping = {
        ".pdf": "research_paper",
        ".docx": "protocol",
        ".csv": "assay_dataset",
        ".xlsx": "assay_dataset",
        ".json": "eln_export",
        ".txt": "regulatory_document",
    }
    return mapping.get(ext, "file_upload")
